#!/usr/bin/env python3
"""
Scraper de df.cl → Word diario.
- Solo noticias del día de hoy
- Incluye las 3 noticias destacadas del tope de cada sección
- Sin duplicados entre secciones
- Paywall: usa archive.today/latest (= Option 1 de RemovePaywalls)
- Al encender el PC recupera automáticamente los días perdidos
"""

import requests
import re
import os
import argparse
from datetime import datetime, timedelta
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

MESES_ES = [
    "enero","febrero","marzo","abril","mayo","junio",
    "julio","agosto","septiembre","octubre","noviembre","diciembre"
]

def fecha_a_texto(fecha: datetime) -> str:
    """Devuelve algo como '10 de abril de 2026'."""
    return f"{fecha.day} de {MESES_ES[fecha.month - 1]} de {fecha.year}"

# ─── CONFIGURACIÓN ────────────────────────────────────────────────────────────

CARPETA_SALIDA = os.environ.get(
    "CARPETA_SALIDA",
    os.path.join(os.path.expanduser("~"), "Desktop", "Noticias_DF"),
)

SECCIONES = [
    {"nombre": "Mercados",      "url": "https://www.df.cl/mercados"},
    {"nombre": "Empresas",      "url": "https://www.df.cl/empresas"},
    {"nombre": "Economía",      "url": "https://www.df.cl/economia_y_politica"},
    {"nombre": "Internacional", "url": "https://www.df.cl/internacional"},
    {"nombre": "Innovación",    "url": "https://www.df.cl/dflab"},
    {"nombre": "Opinión",       "url": "https://www.df.cl/opinion"},
]

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "es-CL,es;q=0.9",
}

MAX_NOTICIAS_POR_SECCION = 20
PAUSA_ENTRE_REQUESTS = 1.5

FRASES_BASURA = {
    "Click acá para ir directamente al contenido",
    "TE PUEDE INTERESAR", "NOTICIAS DESTACADAS",
    "BRANDED CONTENT", "TE RECOMENDAMOS",
    "ARTICULOS RELACIONADOS", "Comparte",
}

# ─── UTILIDADES ───────────────────────────────────────────────────────────────

def fecha_str(fecha: datetime) -> str:
    return fecha.strftime("%d/%m/%Y")

def limpiar_texto(texto: str) -> str:
    texto = re.sub(r'\{\{[^}]+\}\}', '', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    for frase in FRASES_BASURA:
        texto = texto.replace(frase, "")
    return texto.strip()

# ─── LIMPIEZA DE SECCIONES BASURA ─────────────────────────────────────────────

def limpiar_secciones_basura(soup: BeautifulSoup) -> None:
    clases_eliminar = [
        "box__arts", "box__carousel", "cont-banner",
        "vitrina", "branded", "block-tabs",
        "relacionados", "related",
    ]
    for clase in clases_eliminar:
        for el in soup.find_all(class_=re.compile(clase, re.I)):
            el.decompose()
    for tag in soup.find_all(["section", "div", "aside"]):
        texto = (tag.get_text(strip=True) or "")[:60]
        if any(f in texto for f in ["BRANDED CONTENT", "VITRINA EMPRESARIAL",
                                     "PAPEL DIGITAL", "DF MÁS", "INSPÍRATE"]):
            tag.decompose()

# ─── EXTRACCIÓN DE TEXTO ──────────────────────────────────────────────────────

def extraer_texto_articulo(soup: BeautifulSoup) -> str:
    limpiar_secciones_basura(soup)

    # Estrategia 1: archive.today via H3 subtítulos
    h3_subtitulos = [
        h for h in soup.find_all("h3")
        if h.get_text(strip=True) and
           h.get_text(strip=True) not in FRASES_BASURA and
           len(h.get_text(strip=True)) > 5
    ]
    if h3_subtitulos:
        h3 = h3_subtitulos[0]
        try:
            contenedor = h3.parent.parent
            partes = []
            for hijo in contenedor.children:
                if not hasattr(hijo, 'get_text'):
                    continue
                texto_hijo = hijo.get_text(separator=" ", strip=True)
                if any(f in texto_hijo for f in FRASES_BASURA):
                    break
                if texto_hijo and len(texto_hijo) > 20:
                    for h in hijo.find_all("h3"):
                        titulo = h.get_text(strip=True)
                        if titulo and titulo not in FRASES_BASURA:
                            h.replace_with(f"\n\n=== {titulo} ===\n")
                    partes.append(hijo.get_text(separator="\n", strip=True))
            texto = "\n\n".join(partes)
            texto = limpiar_texto(texto)
            if len(texto) > 300:
                return texto
        except Exception:
            pass

    # Estrategia 2: selectores CSS
    for selector in ["#CONTENT", "#SOLID", "article", ".article-body",
                      ".article-content", ".content-body", "main"]:
        el = soup.select_one(selector)
        if el:
            for tag in el.find_all(["script", "style", "nav", "header",
                                     "footer", "aside", "iframe", "button"]):
                tag.decompose()
            for h in el.find_all(["h2", "h3"]):
                titulo = h.get_text(strip=True)
                if titulo and titulo not in FRASES_BASURA:
                    h.replace_with(f"\n\n=== {titulo} ===\n")
            texto = el.get_text(separator="\n", strip=True)
            texto = limpiar_texto(texto)
            if len(texto) > 200:
                return texto

    # Fallback
    parrafos = soup.find_all("p")
    texto = "\n".join(p.get_text(strip=True) for p in parrafos
                      if len(p.get_text(strip=True)) > 50)
    return limpiar_texto(texto) if texto else ""

# ─── PAYWALL ──────────────────────────────────────────────────────────────────

PALABRAS_PAYWALL = [
    "suscr", "subscri", "premium", "regístrate", "inicia sesión",
    "solo para suscriptores", "acceso exclusivo", "contenido exclusivo",
]

def tiene_paywall(soup: BeautifulSoup, html: str) -> bool:
    html_lower = html.lower()
    for palabra in PALABRAS_PAYWALL:
        if palabra in html_lower:
            parrafos = soup.find_all("p")
            if len(" ".join(p.get_text() for p in parrafos)) < 500:
                return True
    for sel in ["[class*='paywall']", "[class*='subscribe']", "[id*='paywall']"]:
        if soup.select_one(sel):
            return True
    return False

def buscar_en_archive(url_original: str) -> tuple:
    """Usa archive.today/latest = Option 1 de RemovePaywalls."""
    try:
        resp = requests.get(
            f"https://archive.today/latest/{url_original}",
            headers=HEADERS, timeout=20, allow_redirects=True
        )
        if "No hay resultados" in resp.text or "no results" in resp.text.lower():
            return "", ""
        soup = BeautifulSoup(resp.text, "html.parser")
        texto = extraer_texto_articulo(soup)
        if texto and len(texto) > 200:
            return texto, resp.url
    except Exception:
        pass
    return "", ""

def extraer_fecha_articulo(soup: BeautifulSoup) -> str:
    """
    Extrae la fecha de publicación del HTML del artículo.
    Retorna algo como '10 de abril de 2026', o '' si no la encuentra.
    """
    # Buscar en meta tags primero (más confiable)
    for meta in soup.find_all("meta"):
        for attr in ["property", "name"]:
            val = meta.get(attr, "")
            if val in ("article:published_time", "datePublished", "pubdate"):
                content = meta.get("content", "")
                try:
                    dt = datetime.fromisoformat(content[:10])
                    return fecha_a_texto(dt)
                except Exception:
                    pass

    # Buscar en el HTML por patrones de fecha en español
    texto_pagina = soup.get_text(" ", strip=True)
    patron = r'\b(\d{1,2})\s+de\s+(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+de\s+(\d{4})\b'
    m = re.search(patron, texto_pagina, re.IGNORECASE)
    if m:
        return f"{int(m.group(1))} de {m.group(2).lower()} de {m.group(3)}"
    return ""


def obtener_contenido_articulo(url: str) -> dict:
    resultado = {"texto": "", "fuente": "df", "tiene_paywall": False, "url_archive": None, "fecha_publicacion": ""}
    try:
        resp = requests.get(url, headers=HEADERS, timeout=12)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        resultado["fecha_publicacion"] = extraer_fecha_articulo(soup)
        texto_directo = extraer_texto_articulo(soup)
        paywall = tiene_paywall(soup, resp.text)

        if not paywall and texto_directo and len(texto_directo) > 300:
            resultado["texto"] = texto_directo
            return resultado

        resultado["tiene_paywall"] = True
        print(f"      🔒 Paywall → archive.today...", end=" ", flush=True)
        texto_archive, url_archive = buscar_en_archive(url)
        if texto_archive:
            resultado["texto"] = texto_archive
            resultado["fuente"] = "archive"
            resultado["url_archive"] = url_archive
            print("✅")
        else:
            resultado["texto"] = texto_directo or ""
            print("❌")
    except Exception:
        pass
    return resultado

# ─── SCRAPING DE SECCIONES ────────────────────────────────────────────────────

def noticia_base(titulo: str, href: str, seccion: str, destacada: bool = False) -> dict:
    return {
        "titulo": titulo, "url": href, "seccion": seccion,
        "contenido": "", "fuente": "df", "tiene_paywall": False,
        "url_archive": None, "destacada": destacada,
    }

def obtener_destacadas(soup: BeautifulSoup, seccion_nombre: str,
                        urls_vistas: set) -> list:
    """
    Obtiene las 3 noticias del tope de cada sección:
    - card__bg = la noticia grande con imagen
    - card__vm = las 2 noticias del lado
    No tienen fecha en el HTML pero siempre son las más recientes.
    """
    destacadas = []
    # Primero la grande, luego las 2 del lado
    selectores = [
        ("card__bg", "desktop"),   # noticia grande
        ("card__vm", None),        # noticias del lado
    ]
    for clase_base, clase_extra in selectores:
        cards = soup.find_all(class_=re.compile(rf'\b{clase_base}\b'))
        for card in cards:
            # Saltar las de "nov" (branded) y las auxi
            clases = card.get("class", [])
            if "card__bg--nov" in clases or "card__auxi" in clases:
                continue
            if clase_extra and clase_extra not in clases:
                continue

            # Buscar link de df.cl
            link = None
            for a in card.find_all("a", href=True):
                href = a["href"]
                if href.startswith("/"):
                    href = "https://www.df.cl" + href
                if "df.cl" in href and "#" not in href and "javascript" not in href:
                    link = href
                    break
            if not link or link in urls_vistas:
                continue

            # Título más largo del card
            titulo = ""
            for tag in card.find_all(["h1", "h2", "h3"]):
                t = tag.get_text(strip=True)
                if len(t) > len(titulo) and len(t) > 15:
                    titulo = t
            if not titulo or len(titulo) < 15:
                continue

            urls_vistas.add(link)
            destacadas.append(noticia_base(titulo, link, seccion_nombre, destacada=True))

            if len(destacadas) >= 3:
                return destacadas

    return destacadas


def obtener_noticias_seccion(seccion: dict, fecha: datetime,
                              urls_vistas: set) -> list:
    """
    Scrapea una sección:
    1. Primero las 3 noticias destacadas del tope (sin filtro de fecha)
    2. Luego las noticias del día con .card__date
    Sin duplicados globales entre secciones.
    """
    fecha_hoy = fecha_str(fecha)
    noticias = []

    try:
        resp = requests.get(seccion["url"], headers=HEADERS, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")

        # 1. Noticias destacadas del tope
        destacadas = obtener_destacadas(soup, seccion["nombre"], urls_vistas)
        noticias.extend(destacadas)

        # 2. Noticias del día con card__date
        for card in soup.find_all(class_=re.compile(r'\bcard\b')):
            fecha_el = card.find(class_=re.compile(r'card__date'))
            if not fecha_el or fecha_el.get_text(strip=True) != fecha_hoy:
                continue

            link = card.find("a", href=True)
            if not link:
                continue
            href = link["href"]
            if href.startswith("/"):
                href = "https://www.df.cl" + href
            if "df.cl" not in href or href in urls_vistas:
                continue
            if any(x in href for x in ["#", "javascript", "mailto"]):
                continue

            titulo = ""
            for tag in card.find_all(["h1", "h2", "h3", "a"]):
                t = tag.get_text(strip=True)
                if len(t) > len(titulo) and len(t) > 15:
                    titulo = t
            if not titulo or len(titulo) < 15:
                continue

            urls_vistas.add(href)
            noticias.append(noticia_base(titulo, href, seccion["nombre"]))

            if len(noticias) >= MAX_NOTICIAS_POR_SECCION:
                break

    except Exception as e:
        print(f"  ⚠️  Error en '{seccion['nombre']}': {e}")

    return noticias

# ─── CREACIÓN DEL WORD ────────────────────────────────────────────────────────

def crear_documento_word(noticias_por_seccion: dict, fecha: datetime,
                          stats: dict) -> str:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Portada
    t = doc.add_heading(level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("📰 Resumen Diario Financiero")
    r.font.color.rgb = RGBColor(0xE6, 0x50, 0x00)
    r.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        fecha.strftime("Edición del %A %d de %B de %Y").title()
    ).font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(
        f"Noticias: {stats['total']}  |  "
        f"Con contenido: {stats['con_contenido']}  |  "
        f"Via archive.today: {stats['de_archive']}"
    ).font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    p = doc.add_paragraph()
    p.add_run("─" * 80).font.color.rgb = RGBColor(0xE6, 0x50, 0x00)
    doc.add_page_break()

    # Noticias por sección
    for seccion, noticias in noticias_por_seccion.items():
        if not noticias:
            continue

        h = doc.add_heading(f"🔹 {seccion}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0xE6, 0x50, 0x00)

        for i, n in enumerate(noticias, 1):
            # Título
            p_t = doc.add_paragraph()
            nr = p_t.add_run(f"{i}. ")
            nr.bold = True
            nr.font.color.rgb = RGBColor(0xE6, 0x50, 0x00)
            tr = p_t.add_run(n["titulo"])
            tr.bold = True
            tr.font.size = Pt(11)

            # Badges
            if n.get("destacada"):
                b = p_t.add_run("  [⭐ Destacada]")
                b.font.size = Pt(9)
                b.font.color.rgb = RGBColor(0xCC, 0x77, 0x00)
            if n["tiene_paywall"] and n["fuente"] == "archive":
                b = p_t.add_run("  [📦 archive.today]")
                b.font.size = Pt(9)
                b.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif n["tiene_paywall"]:
                b = p_t.add_run("  [🔒 paywall]")
                b.font.size = Pt(9)
                b.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

            # URL
            p_u = doc.add_paragraph()
            ur = p_u.add_run(n["url"])
            ur.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
            ur.font.size = Pt(9)
            p_u.paragraph_format.space_before = Pt(0)

            # URL archive
            if n.get("url_archive"):
                p_a = doc.add_paragraph()
                p_a.add_run(f"Archive: {n['url_archive']}").font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                p_a.runs[0].font.size = Pt(9)
                p_a.paragraph_format.space_before = Pt(0)

            # Contenido con subtítulos naranjos
            if n.get("contenido") and len(n["contenido"]) > 50:
                for linea in n["contenido"].split("\n"):
                    linea = linea.strip()
                    if not linea:
                        continue
                    match = re.match(r'^===\s*(.+?)\s*===$', linea)
                    if match:
                        ps = doc.add_paragraph()
                        ps.paragraph_format.left_indent = Inches(0.25)
                        ps.paragraph_format.space_before = Pt(6)
                        sr2 = ps.add_run(match.group(1))
                        sr2.bold = True
                        sr2.font.size = Pt(13)
                        sr2.font.color.rgb = RGBColor(0xE6, 0x50, 0x00)
                    else:
                        pl = doc.add_paragraph()
                        pl.paragraph_format.left_indent = Inches(0.25)
                        pl.paragraph_format.space_before = Pt(1)
                        rl = pl.add_run(linea)
                        rl.font.size = Pt(10)
                        rl.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(2)
            sep.paragraph_format.space_after = Pt(6)
            sep.add_run("· · ·").font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        doc.add_paragraph()

    doc.add_page_break()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.add_run(
        f"Generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}\n"
        f"Fuente: www.df.cl  |  Paywall: archive.today (Option 1 RemovePaywalls)"
    ).font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    carpeta_dia = os.path.join(CARPETA_SALIDA, fecha.strftime("%Y-%m-%d"))
    os.makedirs(carpeta_dia, exist_ok=True)
    ruta = os.path.join(carpeta_dia, fecha.strftime("%Y-%m-%d") + "_DF.docx")
    doc.save(ruta)
    return ruta


# ─── RECUPERAR DÍAS PERDIDOS ──────────────────────────────────────────────────

def detectar_dias_faltantes() -> list:
    """
    Revisa la carpeta de salida y detecta qué días hábiles (lun-vie)
    de los últimos 30 días faltan. Retorna lista de datetimes a scrapear.
    """
    hoy = datetime.now()
    faltantes = []

    for dias_atras in range(1, 31):  # últimos 30 días
        fecha = hoy - timedelta(days=dias_atras)
        # Solo lunes a viernes
        if fecha.weekday() >= 5:
            continue
        nombre = fecha.strftime("%Y-%m-%d") + "_DF.docx"
        ruta = os.path.join(CARPETA_SALIDA, fecha.strftime("%Y-%m-%d"), nombre)
        if not os.path.exists(ruta):
            faltantes.append(fecha)

    return sorted(faltantes)  # del más antiguo al más reciente


# ─── FLUJO PRINCIPAL ──────────────────────────────────────────────────────────

def scrapear_dia(fecha: datetime) -> str:
    print(f"\n📅 Scrapeando: {fecha.strftime('%d/%m/%Y')}")
    noticias_por_seccion = {}
    stats = {"total": 0, "con_contenido": 0, "de_archive": 0}
    urls_vistas = set()

    for seccion in SECCIONES:
        print(f"  🔍 {seccion['nombre']}...", end=" ", flush=True)
        noticias = obtener_noticias_seccion(seccion, fecha, urls_vistas)
        print(f"{len(noticias)} noticias")

        fecha_texto = fecha_a_texto(fecha)  # ej: "10 de abril de 2026"
        noticias_filtradas = []
        for noticia in noticias:
            time.sleep(PAUSA_ENTRE_REQUESTS)
            resultado = obtener_contenido_articulo(noticia["url"])

            # Filtrar destacadas de otro día
            if noticia.get("destacada"):
                fecha_art = resultado.get("fecha_publicacion", "")
                if fecha_art and fecha_art != fecha_texto:
                    print(f"      ⏭️  Destacada excluida (publicada el {fecha_art}): {noticia['titulo'][:60]}")
                    continue

            noticia["contenido"] = resultado["texto"]
            noticia["fuente"] = resultado["fuente"]
            noticia["tiene_paywall"] = resultado["tiene_paywall"]
            noticia["url_archive"] = resultado["url_archive"]

            stats["total"] += 1
            if resultado["texto"]:
                stats["con_contenido"] += 1
            if resultado["fuente"] == "archive":
                stats["de_archive"] += 1

            noticias_filtradas.append(noticia)

        noticias_por_seccion[seccion["nombre"]] = noticias_filtradas
        time.sleep(1)

    ruta_word = crear_documento_word(noticias_por_seccion, fecha, stats)
    carpeta_dia = os.path.join(CARPETA_SALIDA, fecha.strftime("%Y-%m-%d"))
    print(f"\n✅ Guardado en: {carpeta_dia}")
    print(f"   📄 Word: {os.path.basename(ruta_word)}")
    print(f"   Total: {stats['total']}  |  Contenido: {stats['con_contenido']}  |  Archive: {stats['de_archive']}")
    return ruta_word


def main():
    parser = argparse.ArgumentParser(description="Scraper df.cl")
    parser.add_argument("--modo", choices=["dia", "semana", "recuperar"], default="dia",
                        help="dia=hoy | semana=últimos 7 días | recuperar=días perdidos")
    args = parser.parse_args()

    hoy = datetime.now()

    if args.modo == "dia":
        # Solo scrapea HOY — df.cl solo muestra noticias del día actual
        scrapear_dia(hoy)

    elif args.modo == "semana":
        print("📆 Modo semana: últimos 7 días hábiles")
        dias_habiles = [
            hoy - timedelta(days=i)
            for i in range(6, -1, -1)
            if (hoy - timedelta(days=i)).weekday() < 5
        ]
        for i, fecha in enumerate(dias_habiles):
            scrapear_dia(fecha)
            if i < len(dias_habiles) - 1:
                time.sleep(3)

    elif args.modo == "recuperar":
        faltantes = detectar_dias_faltantes()
        if faltantes:
            print(f"\n📋 Recuperando {len(faltantes)} días perdidos...")
            for fecha in faltantes:
                scrapear_dia(fecha)
                time.sleep(3)
        else:
            print("✅ No hay días pendientes.")

    print(f"\n📁 Archivos en: {CARPETA_SALIDA}\n")


if __name__ == "__main__":
    main()
